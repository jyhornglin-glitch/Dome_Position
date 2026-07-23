#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
import_card_hints.py
Parse 小卡關鍵字提示.docx and generate action_hints_data.js.
Supports both old and new numbering in the Word document.
"""

import os
import sys
import json
import re
import docx

DOCX_FILE = "小卡關鍵字提示.docx"
OUTPUT_JS = "card_hints_data.js"

# Map Word location names to formation keys in app.js
# Supports both old and new numbering (shifted by 1)
CATEGORY_MAPPING = {
    # Basic
    '基本': 'basic',
    # Circle (圓形)
    '01圓形': 'circle',
    '02圓形': 'circle',
    # Xing Yuan (行願)
    '02行願': 'xingYuan',
    '03行願': 'xingYuan',
    # Mi Luo (米籮)
    '03米籮': 'miLuo',
    '04米籮': 'miLuo',
    # Jing Si (靜思家風)
    '04靜思家風': 'jingSi',
    '05靜思家風': 'jingSi',
    # Lamp (點一盞燈)
    '05-1有法船(點一盞燈)': 'lamp',
    '06-1有法船(點一盞燈)': 'lamp',
    # No Boat (無法船 - 菜市場5毛錢)
    '05-2無法船(菜市場5毛錢)': 'noBoat',
    '06-2無法船(菜市場5毛錢)': 'noBoat',
    # No Boat 3 (有法船 - 是諸眾生)
    '05-3有法船(是諸眾生)': 'noBoat3',
    '06-3有法船(是諸眾生)': 'noBoat3',
    # Big V (四弘誓願)
    '06四弘誓願': 'bigV',
    '07四弘誓願': 'bigV',
    # Da Chuan Shi (大船師)
    '07-1大船師': 'daChuanShi',
    '08-1大船師': 'daChuanShi',
    # Bone Donation (骨捐能捨)
    '07-2骨捐能捨': 'boneDonation',
    '08-2骨捐能捨': 'boneDonation',
    # Edu (教育)
    '08教育': 'edu',
    '09教育': 'edu',
    # Humanities 1 (人文 09-1 / 10-1)
    '09-1人文': 'humanities1',
    '10-1人文': 'humanities1',
    # Humanities 2 (人文 09-2 / 10-2)
    '09-2人文': 'humanities2',
    '10-2人文': 'humanities2',
    # Five Continents 1 (五大洲 10-1 / 11-1)
    '10-1五大洲': 'fiveContinents1',
    '11-1五大洲': 'fiveContinents1',
    # Five Continents 2 (五大洲 10-2 / 11-2)
    '10-2五大洲': 'fiveContinents2',
    '11-2五大洲': 'fiveContinents2',
    # Flying Apsaras (飛天 11 / 12)
    '11飛天': 'flyingApsaras',
    '12飛天': 'flyingApsaras'
}

def main():
    if not os.path.exists(DOCX_FILE):
        print(f"Error: {DOCX_FILE} not found!")
        sys.exit(1)

    print(f"Parsing DOCX file: {DOCX_FILE}")
    doc = docx.Document(DOCX_FILE)
    table = doc.tables[0]
    R = len(table.rows)
    C = len(table.columns)
    print(f"Opened table with {R} rows, {C} columns.")

    action_hints_data = {
        'basic': [],
        'circle': [],
        'xingYuan': [],
        'miLuo': [],
        'jingSi': [],
        'lamp': [],
        'noBoat': [],
        'noBoat3': [],
        'bigV': [],
        'daChuanShi': [],
        'boneDonation': [],
        'edu': [],
        'humanities1': [],
        'humanities2': [],
        'fiveContinents1': [],
        'fiveContinents2': [],
        'flyingApsaras': []
    }

    for r in range(1, R):
        loc_cell = table.rows[r].cells[0]
        # Clean whitespaces and newlines in location key
        loc_text = loc_cell.text.strip()
        loc_clean = re.sub(r'\s+', '', loc_text) if loc_text else ""
        
        # Normalize parentheses
        loc_clean = loc_clean.replace('（', '(').replace('）', ')')
        
        cat = CATEGORY_MAPPING.get(loc_clean)
        if not cat:
            print(f"Warning: Location '{loc_clean}' not mapped. Skipping.")
            continue
            
        content_cell = table.rows[r].cells[1]
        
        current_item = None
        
        for p in content_cell.paragraphs:
            text = p.text.strip()
            if not text:
                continue
                
            # Check if this paragraph starts a new item section (e.g. starts with 【...】)
            match = re.match(r'^【([^】]+)】(.*)', text)
            if match:
                title_label = match.group(1).strip()
                extra_text = match.group(2).strip()
                
                title = f"【{title_label}】"
                if extra_text.startswith('：') or extra_text.startswith(':'):
                    title += extra_text
                elif extra_text:
                    title += " " + extra_text
                
                current_item = {
                    "title": title,
                    "details": []
                }
                action_hints_data[cat].append(current_item)
            else:
                # If no item created yet, create a default one using the location's name as title
                if not current_item:
                    current_item = {
                        "title": f"【{loc_clean}】",
                        "details": []
                    }
                    action_hints_data[cat].append(current_item)
                
                # Split line breaks in paragraph and append
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                for line in lines:
                    current_item["details"].append({
                        "type": "text",
                        "content": line
                    })

    # Save to card_hints_data.js
    js_content = (
        "// Pocket Slip Card Hints Database — 自動由 import_card_hints.py 產生，請勿手動修改\n"
        f"const CARD_HINTS_DATA = {json.dumps(action_hints_data, ensure_ascii=False, indent=2)};\n\n"
        "// Export if in node environment, otherwise make it global\n"
        "if (typeof module !== 'undefined' && module.exports) {\n"
        "  module.exports = CARD_HINTS_DATA;\n"
        "}\n"
    )

    with open(OUTPUT_JS, 'w', encoding='utf-8') as js_f:
        js_f.write(js_content)

    print(f"Successfully processed DOCX: {DOCX_FILE}")
    print(f"Generated {OUTPUT_JS} successfully!")

if __name__ == "__main__":
    main()
